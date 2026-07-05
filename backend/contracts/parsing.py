"""Extraction of clean text from PDF (pdfplumber) and DOCX (python-docx)."""
import pdfplumber
from docx import Document as DocxDocument


def extract_text(fileobj, file_type: str) -> tuple[str, int]:
    """
    Returns (text, page_count).
    `fileobj` is a file-like object (e.g. io.BytesIO with the file's bytes).
    """
    if file_type == "pdf":
        return _extract_pdf(fileobj)
    if file_type == "docx":
        return _extract_docx(fileobj)
    raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(fileobj) -> tuple[str, int]:
    pages = []
    with pdfplumber.open(fileobj) as pdf:
        for page in pdf.pages:
            # extract_text() reads the text layer; for a scanned PDF it returns None → "".
            pages.append((page.extract_text() or "").strip())
    text = "\n\n".join(pages).strip()
    return text, len(pages)


def _extract_docx(fileobj) -> tuple[str, int]:
    doc = DocxDocument(fileobj)

    lines = [p.text for p in doc.paragraphs if p.text.strip()]

    # Contracts often keep key clauses in tables — so we include those too.
    # Note: python-docx doesn't preserve the paragraph/table order, so tables
    # come last. For contracts that's an acceptable trade-off; rewrite if needed.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    # DOCX has no reliable notion of "pages" → we treat the document as a single block.
    return text, 1
